from copy import deepcopy
from pathlib import Path

from docx import Document


SOURCE = Path("/Users/ahnar/Documents/ChatGPT/论文agent开发/工作周报_2026年8月10日-8月16日.docx")
OUTPUT = Path("/Users/ahnar/Documents/ChatGPT/论文agent开发/工作周报_2026年8月10日-8月16日_含论文研读.docx")

ADDITIONS = [
    "研读《多传感器融合导航：国内技术进展与应用综述（2023）》，整理中文材料并梳理多传感器融合导航技术路线。",
    "研读《GNSS 拒止环境下的无人机导航：计算复杂度、传感器融合与定位方法》，完成全文提取、中文翻译及版面检查，形成 41 页中文译文 PDF。",
]


doc = Document(SOURCE)
heading_two = next(p for p in doc.paragraphs if p.text == "二、本周认识与待改进事项")
bullet_model = next(p for p in doc.paragraphs if p.style.name == "List Bullet")

for text in ADDITIONS:
    new_p = deepcopy(bullet_model._element)
    for node in list(new_p.iter()):
        if node.tag.endswith("}t"):
            node.text = ""
    text_nodes = [node for node in new_p.iter() if node.tag.endswith("}t")]
    if not text_nodes:
        raise RuntimeError("Bullet template has no text node")
    text_nodes[0].text = text
    heading_two._element.addprevious(new_p)

doc.save(OUTPUT)
print(OUTPUT)
