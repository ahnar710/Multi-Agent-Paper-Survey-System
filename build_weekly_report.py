from copy import deepcopy
from pathlib import Path

from docx import Document


REFERENCE = Path("/Users/ahnar/Documents/周报/阿合娜尔周报8.3-8.9/工作周报_2026年8月3日-8月9日.docx")
OUTPUT = Path("/Users/ahnar/Documents/ChatGPT/论文agent开发/工作周报_2026年8月10日-8月16日.docx")

CONTENT = [
    "工作周报",
    "2026年8月10日—8月16日",
    "一、本周工作完成情况",
    "明确 GNSS 文献调研多 Agent 系统的业务目标与验收口径，以每日产出 100 篇可追溯、经验证的全文研究卡片为目标。",
    "完成 MVP 端到端工作流程设计，划分需求规划、检索、初筛、深读、验证、综合和产品分析等 8 类 Agent 职责。",
    "完成 GNSS 多径前沿研究检索计划 v0.1，明确 multipath 与 NLOS 边界、核心检索词、数据源优先级和首轮校准文献方向。",
    "完成主题配置，将日目标分解为信号处理、载波相位建模、机器学习、直接定位、天线阵列和多传感器辅助 6 条技术路线。",
    "二、本周认识与待改进事项",
    "文献产量应按通过全文、引用定位、完整性、质量和去重校验的研究卡片计数，不能以 Agent 调用次数代替有效产出。",
    "当前仍处于方案设计阶段，Harness 工程骨架、实际检索结果和吞吐量测试尚未落地，100 篇/日目标需通过试跑验证。",
    "三、下周工作计划",
    "实现 Harness 最小工作流程骨架，先跑通“主题配置—检索—去重—研究卡片—验证”的单批闭环。",
    "冻结研究卡片 Schema 与校验规则，补齐原文定位、质量评分、产品化判断和失败原因字段。",
    "使用 20—30 篇校准文献进行小规模试跑，测量检索命中率、全文获取率、验证通过率、单篇耗时和成本。",
]


def replace_text_keep_first_run(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


doc = Document(REFERENCE)
paras = doc.paragraphs
if len(paras) != len(CONTENT):
    raise RuntimeError(f"Template paragraph count changed: {len(paras)}")

for paragraph, text in zip(paras, CONTENT):
    replace_text_keep_first_run(paragraph, text)

# Preserve the date styling while ensuring the whole replacement uses it.
for run in paras[1].runs:
    run.font.name = "Calibri"

doc.save(OUTPUT)
print(OUTPUT)
